import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH

"""
This python script is an example to for python docx library
It can read images, txt and generate a new word docx file.
In this example, the script can read all photo in a folder, and make then into a photo album( in the commented out section)
The current code read two photos and insert them into a new table as a photo log. 
Consider using for loop to make all photos into a photo log.
"""



# Path to the folder containing photos
photos_folder = r'Y:\capture'

photoA = r'Y:\capture\photo1.png'
photoB = r'Y:\capture\photo2.png'

# Create a list of photo filenames along with their creation dates
photo_files = os.listdir(photos_folder)
photo_info = []
for filename in photo_files:
    if filename.lower().endswith('.png') and not filename.lower().endswith('png.meta'):
        match = re.match(r'(\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2})\.png', filename)
        if match:
            date_str = match.group(1)
            date_obj = datetime.strptime(date_str, '%Y-%m-%d_%H-%M-%S')
            photo_info.append((filename, date_obj))

# Sort photos by creation date
photo_info.sort(key=lambda x: x[1])

# Create a new Word document
doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
left_margin = Inches(1)
right_margin = Inches(1)
top_margin = Inches(1)
bottom_margin = Inches(1)
section.left_margin = left_margin
section.right_margin = right_margin
section.top_margin = top_margin
section.bottom_margin = bottom_margin

# Create tables and merge certain cells to hold photo
table = doc.add_table(3,3)

#table.allow_autofit = True
table.columns[0].width = Inches(0.84)
table.columns[1].width = Inches(0.84)
table.columns[1].width = Inches(5.83)

table.style = 'Table Grid'


a = table.cell(1,0)
b = table.cell(1,1)
A = a.merge(b)
A.text = "Direction Photo Taken:"

c = table.cell(2,0)
d = table.cell(2,1)
B = c.merge(d)
B.text = "Description:"

table.cell(0,0).text = "Date"
table.cell(0,1).text = "Photo No."

e = table.cell(0,2)
f = table.cell(2,2)
C = e.merge(f)


paragraph = table.cell(1,2).paragraphs[0]

paragraph.add_run().add_picture(photoA, width = Inches(5))
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER




doc.add_paragraph('    ')

# Create tables and merge certain cells to hold photo
table = doc.add_table(3,3)

#table.allow_autofit = True
table.columns[0].width = Inches(0.84)
table.columns[1].width = Inches(0.84)
table.columns[1].width = Inches(5.83)

table.style = 'Table Grid'


a = table.cell(1,0)
b = table.cell(1,1)
A = a.merge(b)
A.text = "Direction Photo Taken:"

c = table.cell(2,0)
d = table.cell(2,1)
B = c.merge(d)
B.text = "Description:"

table.cell(0,0).text = "Date"
table.cell(0,1).text = "Photo No."

e = table.cell(0,2)
f = table.cell(2,2)
C = e.merge(f)

paragraph = table.cell(1,2).paragraphs[0]

paragraph.add_run().add_picture(photoB, width = Inches(5))
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


"""
# Add photos to the document
photos_per_page = 2
for i in range(0, len(photo_info), photos_per_page):
    if i > 0:
        doc.add_page_break()
    for j in range(photos_per_page):
        idx = i + j
        if idx < len(photo_info):
            photo_filename, _ = photo_info[idx]
            photo_path = os.path.join(photos_folder, photo_filename)
            img = Image.open(photo_path)
            img.thumbnail((Inches(6), Inches(6)))
            doc.add_picture(photo_path, width=Inches(6))
            doc.add_paragraph(photo_filename)

# Save the document
# doc.save('photo_album_without_meta.docx')
"""

# Another version. Save the document in the same folder as the photos
docx_filename = os.path.join(photos_folder, 'PhotoLog.docx')
doc.save(docx_filename)
print(f"Document saved as '{docx_filename}'")